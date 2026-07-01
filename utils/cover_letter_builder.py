from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_cover_letter_docx(

    cover_letter,

    filename="Cover_Letter.docx"

):

    document = Document()

    section = document.sections[0]

    section.top_margin = Pt(50)

    section.bottom_margin = Pt(50)

    section.left_margin = Pt(55)

    section.right_margin = Pt(55)

    paragraphs = cover_letter.split("\n")

    for line in paragraphs:

        line = line.strip()

        if not line:

            document.add_paragraph()

            continue

        p = document.add_paragraph()

        run = p.add_run(line)

        run.font.name = "Calibri"

        run.font.size = Pt(11)

        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.save(filename)

    return filename