import os
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch


def add_p_border_bottom(paragraph, color_hex="00468C", size="12"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_hyperlink(paragraph, url, text, color="00468C", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
    return hyperlink

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)

def build_resume_docx(
    resume_data,
    filename="ATS_Resume.docx",
    template="ATS Professional"
):
    if isinstance(resume_data, str):
        try:
            resume_data = json.loads(resume_data)
        except Exception:
            resume_data = {"summary": resume_data, "name": "Format Error Profile"}

    document = Document()

    # Layout Parameters
    font_family = "Calibri"
    show_separators = True
    use_icons = False
    line_spacing = 1.15
    space_after_bullet = Pt(2)
    margin_tb, margin_lr = 0.6, 0.7

    if template == "ATS Professional":
        primary_rgb = RGBColor(0, 70, 140)
        primary_hex = "00468C"
        heading_size, body_size, name_size = 13, 11, 22
        line_spacing = 1.10
    elif template == "Modern":
        primary_rgb = RGBColor(41, 98, 255)
        primary_hex = "2962FF"
        heading_size, body_size, name_size = 14, 11, 24
        use_icons = True
        line_spacing = 1.20
        space_after_bullet = Pt(3)
    elif template == "Executive":
        primary_rgb = RGBColor(64, 64, 64)
        primary_hex = "404040"
        heading_size, body_size, name_size = 14, 11, 22
        font_family = "Times New Roman"
        line_spacing = 1.15
    elif template == "Startup":
        primary_rgb = RGBColor(0, 140, 90)
        primary_hex = "008C5A"
        heading_size, body_size, name_size = 14, 11, 23
        line_spacing = 1.20
    elif template == "Minimal":
        primary_rgb = RGBColor(0, 0, 0)
        primary_hex = "000000"
        heading_size, body_size, name_size = 12, 10, 20
        show_separators = False
        margin_tb, margin_lr = 0.5, 0.5
        line_spacing = 1.05

    section = document.sections[0]
    section.top_margin = Inches(margin_tb)
    section.bottom_margin = Inches(margin_tb)
    section.left_margin = Inches(margin_lr)
    section.right_margin = Inches(margin_lr)

    styles = document.styles
    styles["Normal"].font.name = font_family
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)
    styles["Normal"].font.size = Pt(body_size)

    # Header Identity
    name = resume_data.get("name", "Your Name")
    if name:
        title_p = document.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(2)
        run = title_p.add_run(name)
        run.bold = True
        run.font.size = Pt(name_size)
        run.font.color.rgb = RGBColor(30, 30, 30)

    contact_data = resume_data.get("contact", {})
    contact_p = document.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(12)

    elements = []
    if contact_data.get("email"):
        elements.append((f"📧 {contact_data['email']}" if use_icons else contact_data['email'], None))
    if contact_data.get("phone"):
        elements.append((f"📱 {contact_data['phone']}" if use_icons else contact_data['phone'], None))
    if contact_data.get("portfolio"):
        elements.append(("🌐 Portfolio" if use_icons else "Portfolio", contact_data['portfolio']))
    if contact_data.get("linkedin"):
        elements.append(("💼 LinkedIn" if use_icons else "LinkedIn", contact_data['linkedin']))
    if contact_data.get("github"):
        elements.append(("💻 GitHub" if use_icons else "GitHub", contact_data['github']))

    for i, (text_content, link_url) in enumerate(elements):
        if link_url:
            target_url = link_url if link_url.startswith("http") else f"https://{link_url}"
            add_hyperlink(contact_p, target_url, text_content, color=primary_hex)
        else:
            r = contact_p.add_run(text_content)
            r.font.size = Pt(body_size - 1)
        if i < len(elements) - 1:
            sep_run = contact_p.add_run("  |  ")
            sep_run.font.size = Pt(body_size - 1)
            sep_run.font.color.rgb = RGBColor(140, 140, 140)

    # Layout Priority Order Computation
    has_heavy_exp = len(resume_data.get("experience", [])) > 2
    if template == "Startup":
        layout_order = ["summary", "skills", "projects", "experience", "education", "certifications", "achievements", "languages"]
    elif not has_heavy_exp:
        layout_order = ["summary", "skills", "projects", "education", "experience", "certifications", "achievements", "languages"]
    else:
        layout_order = ["summary", "experience", "skills", "projects", "education", "certifications", "achievements", "languages"]

    def render_section_header(title_text):
        document.add_paragraph()
        h_p = document.add_paragraph()
        h_p.paragraph_format.space_before = Pt(8)
        h_p.paragraph_format.space_after = Pt(4)
        h_p.paragraph_format.keep_with_next = True
        run = h_p.add_run(title_text.upper())
        run.bold = True
        run.font.size = Pt(heading_size)
        run.font.color.rgb = primary_rgb
        if show_separators:
            add_p_border_bottom(h_p, color_hex=primary_hex, size="12")

    # Document Structure Builder Elements
    for block in layout_order:
        if block == "summary" and resume_data.get("summary"):
            render_section_header("Professional Summary")
            p = document.add_paragraph(resume_data["summary"])
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_after = Pt(4)

        elif block == "skills" and resume_data.get("skills"):
            render_section_header("Technical Skills")
            skills = resume_data.get("skills", {})
            for category, items in skills.items():
                if not items:
                    continue
                p = document.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = line_spacing
                title = p.add_run(f"{category}: ")
                title.bold = True
                title.font.size = Pt(body_size)
                body = p.add_run(", ".join(items))
                body.font.size = Pt(body_size)

        elif block == "experience" and resume_data.get("experience"):
            render_section_header("Experience")
            for job in resume_data["experience"]:
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.keep_with_next = True

                r_role = p.add_run(f"{job.get('role', '')} ")
                r_role.bold = True
                p.add_run(f"| {job.get('company', '')} ({job.get('location', '')})")

                p.paragraph_format.tab_stops.add_tab_stop(Inches(6.8), alignment=WD_TAB_ALIGNMENT.RIGHT)
                r_date = p.add_run(f"\t{job.get('duration', '')}")
                r_date.font.color.rgb = RGBColor(100, 100, 100)

                for b in job.get("bullets", []):
                    bp = document.add_paragraph(style="List Bullet")
                    bp.paragraph_format.left_indent = Inches(0.2)
                    bp.paragraph_format.space_after = space_after_bullet
                    bp.paragraph_format.line_spacing = line_spacing
                    bp.add_run(b)

        elif block == "projects" and resume_data.get("projects"):
            render_section_header("Projects")
            for proj in resume_data["projects"]:
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.keep_with_next = True

                r_title = p.add_run(proj.get("title", ""))
                r_title.bold = True

                if proj.get("technologies"):
                    r_tech_label = p.add_run("  •  Technologies: ")
                    r_tech_label.bold = True
                    r_tech_list = p.add_run(", ".join(proj.get("technologies")))
                    r_tech_list.font.italic = True
                    r_tech_list.font.size = Pt(body_size - 0.5)

                for b in proj.get("bullets", []):
                    bp = document.add_paragraph(style="List Bullet")
                    bp.paragraph_format.left_indent = Inches(0.2)
                    bp.paragraph_format.space_after = space_after_bullet
                    bp.paragraph_format.line_spacing = line_spacing
                    bp.add_run(b)

        elif block == "education" and resume_data.get("education"):
            render_section_header("Education")
            for edu in resume_data["education"]:
                p = document.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r_deg = p.add_run(f"{edu.get('degree', '')} ")
                r_deg.bold = True
                p.add_run(f"- {edu.get('institution', '')}")

                p.paragraph_format.tab_stops.add_tab_stop(Inches(6.8), alignment=WD_TAB_ALIGNMENT.RIGHT)
                r_dur = p.add_run(f"\t{edu.get('duration', edu.get('year', ''))}")
                r_dur.font.color.rgb = RGBColor(100, 100, 100)

                details = edu.get("details", "") or (f"CGPA: {edu.get('cgpa')}" if edu.get("cgpa") else "")
                if details:
                    dp = document.add_paragraph(details)
                    dp.paragraph_format.left_indent = Inches(0.15)
                    dp.paragraph_format.space_after = Pt(3)
                    dp.runs[0].font.size = Pt(body_size - 0.5)
                    dp.runs[0].font.italic = True

        elif block == "certifications" and resume_data.get("certifications"):
            render_section_header("Certifications")
            for cert in resume_data["certifications"]:
                bp = document.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)
                bp.add_run(cert)

        elif block == "achievements" and resume_data.get("achievements"):
            render_section_header("Achievements")
            for ach in resume_data["achievements"]:
                bp = document.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)
                bp.add_run(ach)

        elif block == "languages" and resume_data.get("languages"):
            render_section_header("Languages")
            p = document.add_paragraph(", ".join(resume_data["languages"]))
            p.paragraph_format.space_after = Pt(4)

    # Footers
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.8), alignment=WD_TAB_ALIGNMENT.RIGHT)

    f_run = footer_para.add_run("Generated by InterviewGPT AI Resume Studio")
    f_run.font.name = font_family
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = RGBColor(140, 140, 140)

    footer_para.add_run("\tPage ")
    p_run = footer_para.add_run()
    p_run.font.size = Pt(8)
    p_run.font.bold = True
    p_run.font.color.rgb = RGBColor(100, 100, 100)
    add_page_number(p_run)

    document.save(filename)
    return filename


# ---------------------------------------------------------------------------
# ReportLab-based PDF generation (direct resume_data -> PDF, no DOCX step)
# ---------------------------------------------------------------------------

def _escape_xml(text):
    """Escape text so it's safe to embed inside a ReportLab Paragraph."""
    if text is None:
        return ""
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _template_colors(template):
    """Return (primary_hex, heading_size, body_size, name_size) per template."""
    if template == "ATS Professional":
        return "#00468C", 13, 10.5, 22
    elif template == "Modern":
        return "#2962FF", 14, 10.5, 24
    elif template == "Executive":
        return "#404040", 14, 10.5, 22
    elif template == "Startup":
        return "#008C5A", 14, 10.5, 23
    elif template == "Minimal":
        return "#000000", 12, 10, 20
    return "#00468C", 13, 10.5, 22


def add_heading(elements, text, styles, spacer_before=10):
    """Append a section heading (with a preceding spacer) to the elements list."""
    elements.append(Spacer(1, spacer_before))
    elements.append(Paragraph(_escape_xml(text).upper(), styles["SectionHeading"]))
    elements.append(Spacer(1, 4))


def add_bullet(elements, text, styles):
    """Append a single bullet point to the elements list."""
    elements.append(Paragraph(f"•&nbsp;&nbsp;{_escape_xml(text)}", styles["ResumeBullet"]))


def add_paragraph(elements, text, styles, style_name="Body"):
    """Append a plain paragraph to the elements list."""
    elements.append(Paragraph(_escape_xml(text), styles[style_name]))


def _build_styles(primary_hex, heading_size, body_size, name_size):
    base = getSampleStyleSheet()
    primary_color = HexColor(primary_hex)

    base.add(
        _para_style(
            "NameStyle", parent=base["Title"], fontSize=name_size,
            leading=name_size + 4, alignment=TA_CENTER, textColor=HexColor("#1E1E1E"),
            spaceAfter=2,
        )
    )
    base.add(
        _para_style(
            "ContactStyle", parent=base["Normal"], fontSize=9.5,
            leading=13, alignment=TA_CENTER, textColor=HexColor("#8C8C8C"),
            spaceAfter=12,
        )
    )
    base.add(
        _para_style(
            "SectionHeading", parent=base["Normal"], fontSize=heading_size,
            leading=heading_size + 2, textColor=primary_color, spaceAfter=4,
            fontName="Helvetica-Bold",
        )
    )
    base.add(
        _para_style(
            "Body", parent=base["Normal"], fontSize=body_size,
            leading=body_size + 3, spaceAfter=4,
        )
    )
    base.add(
        _para_style(
            "ResumeBullet", parent=base["Normal"], fontSize=body_size,
            leading=body_size + 3, leftIndent=14, spaceAfter=2,
        )
    )
    base.add(
        _para_style(
            "SkillLine", parent=base["Normal"], fontSize=body_size,
            leading=body_size + 3, spaceAfter=3,
        )
    )
    base.add(
        _para_style(
            "EntryTitle", parent=base["Normal"], fontSize=body_size + 0.5,
            leading=body_size + 4, spaceBefore=4, spaceAfter=1,
            fontName="Helvetica-Bold",
        )
    )
    base.add(
        _para_style(
            "EntryMeta", parent=base["Normal"], fontSize=body_size - 1,
            leading=body_size + 2, textColor=HexColor("#646464"), spaceAfter=3,
        )
    )
    return base


def _para_style(name, parent, **kwargs):
    from reportlab.lib.styles import ParagraphStyle
    return ParagraphStyle(name, parent=parent, **kwargs)


def _footer_canvas(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(HexColor("#8C8C8C"))
    canvas.drawString(doc.leftMargin, 0.4 * inch, "Generated by InterviewGPT AI Resume Studio")
    canvas.setFillColor(HexColor("#646464"))
    canvas.drawRightString(
        doc.pagesize[0] - doc.rightMargin, 0.4 * inch, f"Page {canvas.getPageNumber()}"
    )
    canvas.restoreState()


def generate_resume_pdf(resume_data, template="ATS Professional", pdf_filename="ATS_Resume.pdf", **kwargs):
    # Extract template argument dynamically if provided via kwargs key mapping
    selected_template = kwargs.get("template", template)

    if isinstance(resume_data, str):
        try:
            resume_data = json.loads(resume_data)
        except Exception:
            resume_data = {"summary": resume_data, "name": "Format Error Profile"}

    primary_hex, heading_size, body_size, name_size = _template_colors(selected_template)
    styles = _build_styles(primary_hex, heading_size, body_size, name_size)

    doc = SimpleDocTemplate(
        pdf_filename,
        pagesize=letter,
        topMargin=0.6 * inch,
        bottomMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
    )

    elements = []

    # --- Header ---
    name = resume_data.get("name", "Your Name")
    if name:
        elements.append(Paragraph(_escape_xml(name), styles["NameStyle"]))

    contact_data = resume_data.get("contact", {})
    contact_bits = []
    if contact_data.get("email"):
        contact_bits.append(_escape_xml(contact_data["email"]))
    if contact_data.get("phone"):
        contact_bits.append(_escape_xml(contact_data["phone"]))
    def _normalize_url(url):
        if not url.startswith(("http://", "https://")):
            url = "https://" + url
        return url

    if contact_data.get("portfolio"):
        url = _normalize_url(contact_data["portfolio"])
        contact_bits.append(f'<link href="{url}">Portfolio</link>')
    if contact_data.get("linkedin"):
        url = _normalize_url(contact_data["linkedin"])
        contact_bits.append(f'<link href="{url}">LinkedIn</link>')
    if contact_data.get("github"):
        url = _normalize_url(contact_data["github"])
        contact_bits.append(f'<link href="{url}">GitHub</link>')
    if contact_bits:
        elements.append(Paragraph("&nbsp;|&nbsp;".join(contact_bits), styles["ContactStyle"]))

    # --- Layout order (mirrors DOCX logic) ---
    has_heavy_exp = len(resume_data.get("experience", [])) > 2
    if selected_template == "Startup":
        layout_order = ["summary", "skills", "projects", "experience", "education", "certifications", "achievements", "languages"]
    elif not has_heavy_exp:
        layout_order = ["summary", "skills", "projects", "education", "experience", "certifications", "achievements", "languages"]
    else:
        layout_order = ["summary", "experience", "skills", "projects", "education", "certifications", "achievements", "languages"]

    for block in layout_order:
        if block == "summary" and resume_data.get("summary"):
            add_heading(elements, "Professional Summary", styles)
            add_paragraph(elements, resume_data["summary"], styles)

        elif block == "skills" and resume_data.get("skills"):
            add_heading(elements, "Technical Skills", styles)
            for category, items in resume_data.get("skills", {}).items():
                if not items:
                    continue
                line = f"<b>{_escape_xml(category)}:</b> {_escape_xml(', '.join(items))}"
                elements.append(Paragraph(line, styles["SkillLine"]))

        elif block == "experience" and resume_data.get("experience"):
            add_heading(elements, "Experience", styles)
            for job in resume_data["experience"]:
                title_line = f"<b>{_escape_xml(job.get('role', ''))}</b> | {_escape_xml(job.get('company', ''))} ({_escape_xml(job.get('location', ''))})"
                elements.append(Paragraph(title_line, styles["EntryTitle"]))
                if job.get("duration"):
                    elements.append(Paragraph(_escape_xml(job["duration"]), styles["EntryMeta"]))
                for b in job.get("bullets", []):
                    add_bullet(elements, b, styles)

        elif block == "projects" and resume_data.get("projects"):
            add_heading(elements, "Projects", styles)
            for proj in resume_data["projects"]:
                title_line = f"<b>{_escape_xml(proj.get('title', ''))}</b>"
                if proj.get("technologies"):
                    title_line += f"&nbsp;&nbsp;•&nbsp;&nbsp;<b>Technologies:</b> <i>{_escape_xml(', '.join(proj['technologies']))}</i>"
                elements.append(Paragraph(title_line, styles["EntryTitle"]))
                for b in proj.get("bullets", []):
                    add_bullet(elements, b, styles)

        elif block == "education" and resume_data.get("education"):
            add_heading(elements, "Education", styles)
            for edu in resume_data["education"]:
                title_line = f"<b>{_escape_xml(edu.get('degree', ''))}</b> - {_escape_xml(edu.get('institution', ''))}"
                elements.append(Paragraph(title_line, styles["EntryTitle"]))
                duration = edu.get("duration", edu.get("year", ""))
                if duration:
                    elements.append(Paragraph(_escape_xml(duration), styles["EntryMeta"]))
                details = edu.get("details", "") or (f"CGPA: {edu.get('cgpa')}" if edu.get("cgpa") else "")
                if details:
                    elements.append(Paragraph(f"<i>{_escape_xml(details)}</i>", styles["EntryMeta"]))

        elif block == "certifications" and resume_data.get("certifications"):
            add_heading(elements, "Certifications", styles)
            for cert in resume_data["certifications"]:
                add_bullet(elements, cert, styles)

        elif block == "achievements" and resume_data.get("achievements"):
            add_heading(elements, "Achievements", styles)
            for ach in resume_data["achievements"]:
                add_bullet(elements, ach, styles)

        elif block == "languages" and resume_data.get("languages"):
            add_heading(elements, "Languages", styles)
            add_paragraph(elements, ", ".join(resume_data["languages"]), styles)

    doc.build(elements, onFirstPage=_footer_canvas, onLaterPages=_footer_canvas)

    return pdf_filename